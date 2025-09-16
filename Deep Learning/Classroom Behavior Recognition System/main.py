import logging
import os
import sys
from datetime import datetime
import sys
import os
# ===== 在文件开头添加 =====
from flask import Flask, render_template, request, jsonify, send_from_directory
if getattr(sys, 'frozen', False):
    base_path = sys._MEIPASS
else:
    base_path = os.path.dirname(os.path.abspath(__file__))

# 修改 Flask 初始化部分
app = Flask(__name__, 
            template_folder=os.path.join(base_path, 'web'),  # 动态模板路径
            static_folder=os.path.join(base_path, 'web'))     # 动态静态文件路径
# ********* 日志配置开始 *********
LOG_DIR = "logs"  # 日志目录
os.makedirs(LOG_DIR, exist_ok=True)  # 创建日志目录（若不存在）

# 生成日志文件名（带时间戳）
log_filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
LOG_PATH = os.path.join(LOG_DIR, log_filename)

# 配置日志格式
logging.basicConfig(
    level=logging.DEBUG,  # 设置日志级别为 DEBUG
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(LOG_PATH, encoding='utf-8'),  # 写入日志文件
        logging.StreamHandler(sys.stdout)  # 同时输出到控制台
    ]
)
import os
import sys
import cv2
import torch
import traceback
import base64
import numpy as np
from flask import Flask, render_template, request, jsonify, send_from_directory
from werkzeug.utils import secure_filename
from ultralytics import YOLO
from PIL import Image, ImageDraw, ImageFont
import io
import json
from datetime import datetime

app = Flask(__name__, template_folder='web', static_folder='web')

# 配置
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['RESULTS_FOLDER'] = 'classroom_behavior_results'

# 确保上传和结果文件夹存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['RESULTS_FOLDER'], exist_ok=True)

class ClassroomBehaviorSystem:
    def __init__(self):
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        self.model = None
        self.detection_results = []
        self.conf_threshold = 0.25
        self.iou_threshold = 0.45
        
        # 行为类别映射
        self.class_names = {
            0: "使用手机", 1: "弯腰", 2: "看书", 3: "低头",
            4: "举手", 5: "手持手机", 6: "抬头", 7: "阅读",
            8: "睡觉", 9: "转头", 10: "端正坐姿", 11: "书写"
        }
        
        print(f"系统初始化完成，使用设备: {self.device}")
    
    def load_model(self, model_path):
        """加载YOLO模型"""
        try:
            self.model = YOLO(model_path).to(self.device)
            print(f"模型加载成功: {model_path}")
            return True, "模型加载成功"
        except Exception as e:
            print(f"模型加载失败: {str(e)}")
            return False, str(e)
    
    def detect_image(self, image_path, conf_threshold=None, iou_threshold=None):
        """检测图片中的行为"""
        if self.model is None:
            return False, "模型未加载", [], None
        
        try:
            # 使用传入的阈值或默认值
            conf = conf_threshold if conf_threshold is not None else self.conf_threshold
            iou = iou_threshold if iou_threshold is not None else self.iou_threshold
            
            # 读取图片
            image = cv2.imread(image_path)
            if image is None:
                return False, "无法读取图片文件", [], None
            
            # 执行检测
            results = self.model.predict(
                image,
                conf=conf,
                iou=iou,
                verbose=False
            )
            
            # 解析检测结果
            detections = self._parse_detections(results)
            
            # 绘制检测框
            annotated_image = self._draw_detection_boxes(image.copy(), detections)
            
            # 将图片转换为base64
            _, buffer = cv2.imencode('.jpg', annotated_image)
            img_base64 = base64.b64encode(buffer).decode('utf-8')
            
            return True, "检测完成", detections, img_base64
            
        except Exception as e:
            print(f"图片检测失败: {str(e)}")
            return False, str(e), [], None
    
    def detect_frame(self, image_data, conf_threshold=None, iou_threshold=None):
        """检测单帧图像"""
        if self.model is None:
            return False, "模型未加载", []
        
        try:
            # 使用传入的阈值或默认值
            conf = conf_threshold if conf_threshold is not None else self.conf_threshold
            iou = iou_threshold if iou_threshold is not None else self.iou_threshold
            
            # 将字节数据转换为图像
            nparr = np.frombuffer(image_data, np.uint8)
            image = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
            
            if image is None:
                return False, "无法解析图像数据", []
            
            # 执行检测
            results = self.model.predict(
                image,
                conf=conf,
                iou=iou,
                verbose=False
            )
            
            # 解析检测结果
            detections = self._parse_detections(results)
            
            return True, "检测完成", detections
            
        except Exception as e:
            print(f"帧检测失败: {str(e)}")
            return False, str(e), []
    
    def _parse_detections(self, results):
        """解析YOLO检测结果"""
        detections = []
        
        for result in results:
            for box in result.boxes:
                class_id = int(box.cls)
                confidence = float(box.conf)
                bbox = list(map(int, box.xyxy[0].tolist()))
                
                detection = {
                    'class_id': class_id,
                    'class_name': self.class_names.get(class_id, "未知"),
                    'confidence': confidence,
                    'bbox': bbox
                }
                detections.append(detection)
        
        return detections
    
    def _draw_detection_boxes(self, frame, detections):
        """绘制检测框和标签"""
        # 先用OpenCV绘制矩形框
        for detection in detections:
            x1, y1, x2, y2 = detection['bbox']
            class_id = detection['class_id']
            confidence = detection['confidence']
            
            # 颜色映射
            colors = {
                0: (50, 50, 255),   # 使用手机 - 红色
                4: (50, 255, 50),   # 举手 - 绿色
                8: (255, 50, 50),   # 睡觉 - 蓝色
                10: (50, 255, 255)  # 端正坐姿 - 黄色
            }
            color = colors.get(class_id, (100, 100, 100))  # 默认灰色
            
            # 绘制矩形框
            cv2.rectangle(frame, (x1, y1), (x2, y2), color, 2)
        
        # 使用PIL绘制中文标签
        pil_img = Image.fromarray(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB))
        draw = ImageDraw.Draw(pil_img)
        
        # 尝试加载中文字体
        try:
            font = ImageFont.truetype("simhei.ttf", 24)
        except:
            try:
                font = ImageFont.truetype("arial.ttf", 20)
            except:
                font = ImageFont.load_default()
        
        # 绘制文本标签
        for detection in detections:
            x1, y1, x2, y2 = detection['bbox']
            label = f"{detection['class_name']} {detection['confidence']:.2f}"
            
            # 计算文本大小
            bbox = draw.textbbox((0, 0), label, font=font)
            text_width = bbox[2] - bbox[0]
            text_height = bbox[3] - bbox[1]
            
            # 绘制文本背景
            draw.rectangle(
                [x1, y1 - text_height - 5, x1 + text_width + 10, y1],
                fill=(0, 0, 0)
            )
            
            # 绘制文本
            draw.text((x1 + 5, y1 - text_height - 5), label, font=font, fill=(255, 255, 255))
        
        # 转换回OpenCV格式
        final_frame = cv2.cvtColor(np.array(pil_img), cv2.COLOR_RGB2BGR)
        return final_frame
    

# ***************************************************************************************************************
# ***************************************************************************************************************
# ***************************************************************************************************************
    def save_results(self, image_path, detections, file_type='image'):
        """保存检测结果，生成Word文档报告"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = os.path.basename(image_path)
            base_filename = f"behavior_{timestamp}_{os.path.splitext(filename)[0]}"
            result_dir = os.path.join(app.config['RESULTS_FOLDER'], base_filename)
            os.makedirs(result_dir, exist_ok=True)
            
            # 读取原始图像
            original_image = cv2.imread(image_path)
            if original_image is None:
                return False, f"无法读取原始图像: {image_path}"
            
            # 绘制检测框得到带标注的图像
            annotated_image = self._draw_detection_boxes(original_image.copy(), detections)
            
            # 保存原始和标注图像
            orig_path = os.path.join(result_dir, "original.jpg")
            anno_path = os.path.join(result_dir, "annotated.jpg")
            cv2.imwrite(orig_path, original_image)
            cv2.imwrite(anno_path, annotated_image)
            
            # 创建Word文档
            doc = Document()
            
            #字体设置
            doc.styles['Normal'].font.name = 'SimSun'  # 西文字体
            doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 中文字体

            # 添加标题
            title = doc.add_paragraph()
            title_run = title.add_run(f"课堂行为检测报告 - {datetime.now().strftime('%Y/%m/%d %H:%M:%S')}")
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加基本信息
            doc.add_paragraph(f"文件名称: {filename}")
            doc.add_paragraph(f"检测时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
            doc.add_paragraph(f"检测到的行为总数: {len(detections)}个")
            
            # 添加原始和标注图像
            doc.add_paragraph("原始图像和检测结果:")
            images = doc.add_paragraph()
            images.alignment = WD_ALIGN_PARAGRAPH.CENTER
            images.add_run().add_picture(orig_path, width=Inches(4))
            images.add_run("      ")  # 添加间距
            images.add_run().add_picture(anno_path, width=Inches(4))
            
            # 添加行为详情
            doc.add_paragraph("行为详情:", style='List Bullet').bold = True
            
            for i, det in enumerate(detections):
                # 添加行为标题
                doc.add_paragraph(f"行为#{i+1}: {det['class_name']} (置信度: {det['confidence']:.2f})", style='Heading 3')
                
                # 添加坐标
                x1, y1, x2, y2 = det['bbox']
                doc.add_paragraph(f"坐标: ({x1}, {y1}) - ({x2}, {y2})")
                
                # 添加行为截图
                crop_img = original_image[y1:y2, x1:x2]
                if crop_img.size > 0:
                    crop_path = os.path.join(result_dir, f"crop_{i}.jpg")
                    cv2.imwrite(crop_path, crop_img)
                    
                    # 添加到文档
                    crop_p = doc.add_paragraph()
                    crop_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    crop_p.add_run().add_picture(crop_path, width=Inches(2))
            
            # 添加统计信息
            behavior_count = {}
            for det in detections:
                behavior = det['class_name']
                behavior_count[behavior] = behavior_count.get(behavior, 0) + 1
            
            doc.add_page_break()
            doc.add_paragraph("行为统计:", style='Heading 2')
            
            # 添加统计表格
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '行为名称'
            hdr_cells[1].text = '出现次数'
            hdr_cells[2].text = '占比(%)'
            
            for behavior, count in behavior_count.items():
                row_cells = table.add_row().cells
                row_cells[0].text = behavior
                row_cells[1].text = str(count)
                row_cells[2].text = f"{count / len(detections) * 100:.1f}"
            
            # 保存Word文档
            doc_path = os.path.join(result_dir, f"{base_filename}_report.docx")
            doc.save(doc_path)
            
            return True, doc_path
            
        except Exception as e:
            print(f"保存结果失败: {str(e)}\n{traceback.format_exc()}")
            return False, str(e)
# 创建系统实例
behavior_system = ClassroomBehaviorSystem()

# 路由定义
@app.route('/')
def index():
    """主页"""
    return render_template('index.html')

@app.route('/api/load_model', methods=['POST'])
def load_model():
    """加载模型API"""
    try:
        if 'model' not in request.files:
            return jsonify({'success': False, 'error': '没有选择模型文件'})
        
        file = request.files['model']
        if file.filename == '':
            return jsonify({'success': False, 'error': '没有选择模型文件'})
        
        if not file.filename.endswith('.pt'):
            return jsonify({'success': False, 'error': '只支持.pt格式的模型文件'})
        
        # 保存模型文件
        filename = secure_filename(file.filename)
        model_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(model_path)
        
        # 加载模型
        success, message = behavior_system.load_model(model_path)
        
        if success:
            return jsonify({
                'success': True,
                'message': message,
                'model_info': {
                    'filename': filename,
                    'device': str(behavior_system.device)
                }
            })
        else:
            return jsonify({'success': False, 'error': message})
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/detect_image', methods=['POST'])
def detect_image():
    """图片检测API"""
    try:
        if 'image' not in request.files:
            return jsonify({'success': False, 'error': '没有选择图片文件'})
        
        file = request.files['image']
        if file.filename == '':
            return jsonify({'success': False, 'error': '没有选择图片文件'})
        
        # 获取参数
        conf_threshold = float(request.form.get('conf_threshold', behavior_system.conf_threshold))
        iou_threshold = float(request.form.get('iou_threshold', behavior_system.iou_threshold))
        save_results = request.form.get('save_results', 'false').lower() == 'true'
        
        # 保存图片文件
        filename = secure_filename(file.filename)
        image_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(image_path)
        
        # 执行检测
        success, message, detections, annotated_image = behavior_system.detect_image(
            image_path, conf_threshold, iou_threshold
        )
        
        if success:
            result = {
                'success': True,
                'message': message,
                'detections': detections,
                'annotated_image': annotated_image
            }
            
            # 保存结果
            if save_results and detections:
                save_success, save_path = behavior_system.save_results(image_path, detections, 'image')
                if save_success:
                    result['saved_to'] = save_path
            
            return jsonify(result)
        else:
            return jsonify({'success': False, 'error': message})
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/detect_video', methods=['POST'])
def detect_video():
    """视频检测API"""
    try:
        if 'video' not in request.files:
            return jsonify({'success': False, 'error': '没有选择视频文件'})
        
        file = request.files['video']
        if file.filename == '':
            return jsonify({'success': False, 'error': '没有选择视频文件'})
        
        # 获取参数
        conf_threshold = float(request.form.get('conf_threshold', behavior_system.conf_threshold))
        iou_threshold = float(request.form.get('iou_threshold', behavior_system.iou_threshold))
        save_results = request.form.get('save_results', 'false').lower() == 'true'
        
        # 保存视频文件
        filename = secure_filename(file.filename)
        video_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(video_path)
        
        # 执行视频检测（这里简化处理，实际可以提取关键帧进行检测）
        import cv2
        cap = cv2.VideoCapture(video_path)
        all_detections = []
        
        # 提取第一帧进行检测作为示例p
        ret, frame = cap.read()
        if ret:
            # 将帧转换为字节数据
            _, buffer = cv2.imencode('.jpg', frame)
            frame_data = buffer.tobytes()
            
            # 执行检测
            success, message, detections = behavior_system.detect_frame(
                frame_data, conf_threshold, iou_threshold
            )
            
            if success:
                all_detections = detections
        
        cap.release()
        
        result = {
            'success': True,
            'message': f'视频检测完成，检测到 {len(all_detections)} 个行为',
            'detections': all_detections
        }
        
        # 保存结果
        if save_results and all_detections:
            save_success, save_path = behavior_system.save_results(video_path, all_detections, 'video')
            if save_success:
                result['saved_to'] = save_path
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/detect_frame', methods=['POST'])
def detect_frame():
    """帧检测API（用于实时检测）"""
    try:
        if 'image' not in request.files:
            return jsonify({'success': False, 'error': '没有图像数据'})
        
        file = request.files['image']
        
        # 获取参数
        conf_threshold = float(request.form.get('conf_threshold', behavior_system.conf_threshold))
        iou_threshold = float(request.form.get('iou_threshold', behavior_system.iou_threshold))
        
        # 读取图像数据
        image_data = file.read()
        
        # 执行检测
        success, message, detections = behavior_system.detect_frame(
            image_data, conf_threshold, iou_threshold
        )
        
        if success:
            return jsonify({
                'success': True,
                'message': message,
                'detections': detections
            })
        else:
            return jsonify({'success': False, 'error': message})
            
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/status')
def get_status():
    """获取系统状态"""
    return jsonify({
        'model_loaded': behavior_system.model is not None,
        'device': str(behavior_system.device),
        'conf_threshold': behavior_system.conf_threshold,
        'iou_threshold': behavior_system.iou_threshold
    })

# 静态文件服务
@app.route('/<path:filename>')
def static_files(filename):
    return send_from_directory('web', filename)

if __name__ == '__main__':
    print("课堂行为检测系统 Web版本")
    app.run(host='0.0.0.0', port=5000, debug=False, threaded=True)